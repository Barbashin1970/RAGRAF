"""Хранилище загруженных аналитиком документов (PDF/DOCX) для контекста Q&A.

Пайплайн при upload:
  bytes → parse_text() → chunk_text() → bge-m3 embed → INSERT в DuckDB

Пайплайн при chat:
  query → embed_query → cosine search по chunks из enabled-документов → top-k

Лимиты:
- Максимум 10 документов на инстанс (single-user, см. ADR-003).
- Размер одного документа — до 10 МБ.
- chunk_size = 800 символов с overlap 120 — компромисс между гранулярностью
  retrieval'а и количеством обращений к bge-m3.

NotebookLM-style: каждый документ имеет toggle enabled (включён ли в контекст);
аналитик отмечает 1-3 источника для конкретного вопроса.
"""
from __future__ import annotations

import json
import re
import threading
import uuid
from datetime import datetime, timezone
from io import BytesIO
from typing import Any

from app.services import regulation_store

_LOCK = threading.RLock()

MAX_DOCUMENTS = 10
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 МБ
CHUNK_SIZE = 800
CHUNK_OVERLAP = 120


# ── Parsing ───────────────────────────────────────────────────────────


def parse_pdf(data: bytes) -> str:
    """Извлечь текст из PDF (pypdf, чистый Python, без системных зависимостей)."""
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(p.strip() for p in pages if p.strip())


def parse_docx(data: bytes) -> str:
    """Извлечь текст из DOCX (python-docx). Берём параграфы и таблицы."""
    from docx import Document

    doc = Document(BytesIO(data))
    blocks: list[str] = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    # Также таблицы — внутри них могут быть нормативные параметры.
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    return "\n\n".join(blocks)


def parse_document(filename: str, mime_type: str, data: bytes) -> str:
    """Диспетчер парсинга по MIME / расширению. Кидает ValueError для неподдерживаемых."""
    lower = filename.lower()
    if mime_type == "application/pdf" or lower.endswith(".pdf"):
        return parse_pdf(data)
    if (
        mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or lower.endswith(".docx")
    ):
        return parse_docx(data)
    raise ValueError(f"Неподдерживаемый формат: {mime_type or filename}. Только PDF/DOCX.")


# ── Chunking ──────────────────────────────────────────────────────────


def chunk_text(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    """Разбить текст на чанки с overlap.

    Стратегия: режем по границам предложений (точка + пробел), если возможно;
    fallback на жёсткий размер. Overlap нужен чтобы факт, разрезанный по
    границе чанка, оставался findable хотя бы из одного контекста.
    """
    text = re.sub(r"\s+", " ", text).strip()
    if not text:
        return []
    if len(text) <= size:
        return [text]

    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + size, len(text))
        # Пытаемся завершить на границе предложения
        if end < len(text):
            for sep in (". ", "! ", "? ", "\n\n"):
                last = text.rfind(sep, start, end)
                if last != -1 and last > start + size // 2:
                    end = last + len(sep)
                    break
        chunks.append(text[start:end].strip())
        if end >= len(text):
            break
        start = end - overlap
    return [c for c in chunks if c]


# ── Embedding ─────────────────────────────────────────────────────────


async def embed_chunks(chunks: list[str]) -> list[list[float]] | None:
    """Эмбеддит чанки через bge-m3 одним батчем. None если Ollama недоступна."""
    if not chunks:
        return []
    from app.config import settings

    if not settings.embeddings_enabled:
        # Embeddings отключены — документ сохранится с embedding=NULL, retrieval
        # деградирует на keyword-match (это уже сделано ниже по коду).
        return None

    try:
        from openai import AsyncOpenAI

        client = AsyncOpenAI(
            base_url=settings.openai_base_url or None,
            api_key=settings.openai_api_key or "ollama",
            timeout=120.0,
        )
        resp = await client.embeddings.create(
            model=settings.ragu_embed_model,
            input=chunks,
        )
        return [list(d.embedding) for d in resp.data]
    except Exception:
        return None


async def embed_query(query: str) -> list[float] | None:
    """Эмбеддит один запрос. None при ошибке Ollama."""
    if not query.strip():
        return None
    from app.config import settings

    if not settings.embeddings_enabled:
        return None

    try:
        from openai import AsyncOpenAI

        client = AsyncOpenAI(
            base_url=settings.openai_base_url or None,
            api_key=settings.openai_api_key or "ollama",
            timeout=60.0,
        )
        resp = await client.embeddings.create(
            model=settings.ragu_embed_model,
            input=query,
        )
        return list(resp.data[0].embedding)
    except Exception:
        return None


def _cosine(a: list[float], b: list[float]) -> float:
    """Cosine similarity без numpy — для малых корпусов (до 1000 чанков) хватает."""
    if not a or not b or len(a) != len(b):
        return 0.0
    import math

    dot = sum(x * y for x, y in zip(a, b))
    na = math.sqrt(sum(x * x for x in a))
    nb = math.sqrt(sum(x * x for x in b))
    if na == 0 or nb == 0:
        return 0.0
    return dot / (na * nb)


# ── Public API ────────────────────────────────────────────────────────


def list_documents() -> list[dict[str, Any]]:
    """Список документов аналитика, отсортированный по дате загрузки (новые сверху)."""
    with _LOCK:
        c = regulation_store._connection()
        rows = c.execute(
            """
            SELECT doc_id, filename, mime_type, size_bytes, uploaded_at,
                   enabled, total_chunks, char_count, error
            FROM user_documents
            ORDER BY uploaded_at DESC
            """
        ).fetchall()
    return [
        {
            "doc_id": r[0],
            "filename": r[1],
            "mime_type": r[2],
            "size_bytes": int(r[3]),
            "uploaded_at": r[4].isoformat() if hasattr(r[4], "isoformat") else str(r[4]),
            "enabled": bool(r[5]),
            "total_chunks": int(r[6]),
            "char_count": int(r[7]),
            "error": r[8],
        }
        for r in rows
    ]


def count_documents() -> int:
    with _LOCK:
        c = regulation_store._connection()
        row = c.execute("SELECT COUNT(*) FROM user_documents").fetchone()
    return int(row[0]) if row else 0


def count_enabled() -> int:
    with _LOCK:
        c = regulation_store._connection()
        row = c.execute("SELECT COUNT(*) FROM user_documents WHERE enabled = TRUE").fetchone()
    return int(row[0]) if row else 0


async def add_document(filename: str, mime_type: str, data: bytes) -> dict[str, Any]:
    """Полный пайплайн загрузки: parse → chunk → embed → store. Возвращает метаданные.

    Если эмбеддинг не удался (Ollama недоступна), документ всё равно сохраняется
    с embedding=NULL — retrieval тогда деградирует на keyword-match (TF-IDF).
    """
    if count_documents() >= MAX_DOCUMENTS:
        raise ValueError(
            f"Достигнут лимит {MAX_DOCUMENTS} документов. Удалите старые перед загрузкой новых."
        )
    if len(data) > MAX_FILE_SIZE:
        raise ValueError(
            f"Размер файла {len(data) / 1024 / 1024:.1f} МБ превышает лимит "
            f"{MAX_FILE_SIZE / 1024 / 1024:.0f} МБ."
        )

    text = parse_document(filename, mime_type, data)
    if not text.strip():
        raise ValueError("Не удалось извлечь текст из документа (возможно, скан без OCR).")

    chunks = chunk_text(text)
    embeddings = await embed_chunks(chunks)
    # embeddings может быть None если Ollama недоступна — храним без векторов

    doc_id = f"doc_{uuid.uuid4().hex[:12]}"
    now = datetime.now(timezone.utc)

    with _LOCK:
        c = regulation_store._connection()
        c.begin()
        try:
            c.execute(
                """
                INSERT INTO user_documents
                    (doc_id, filename, mime_type, size_bytes, uploaded_at,
                     enabled, total_chunks, char_count, error)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                [
                    # enabled=False по дефолту — opt-in модель: пользователь
                    # сам включает галку в DocumentsPanel когда документ нужен
                    # в контексте. Раньше дефолт был True — это раздувало
                    # prompt при первом «Привет» (LLM получала весь корпус +
                    # все загруженные доки сразу). Теперь fresh-chat пустой
                    # до явного выбора пользователя.
                    doc_id, filename, mime_type, len(data), now,
                    False, len(chunks), len(text),
                    None if embeddings else "embeddings_unavailable",
                ],
            )
            for i, chunk in enumerate(chunks):
                vec = embeddings[i] if embeddings else None
                c.execute(
                    """
                    INSERT INTO document_chunks (chunk_id, doc_id, chunk_index, text, embedding)
                    VALUES (?, ?, ?, ?, ?)
                    """,
                    [
                        f"chunk_{uuid.uuid4().hex[:12]}",
                        doc_id, i, chunk,
                        json.dumps(vec) if vec is not None else None,
                    ],
                )
            c.commit()
        except Exception:
            c.rollback()
            raise

    return {
        "doc_id": doc_id,
        "filename": filename,
        "mime_type": mime_type,
        "size_bytes": len(data),
        "uploaded_at": now.isoformat(),
        "enabled": False,
        "total_chunks": len(chunks),
        "char_count": len(text),
        "error": None if embeddings else "embeddings_unavailable",
    }


def toggle_document(doc_id: str, enabled: bool) -> dict[str, Any] | None:
    with _LOCK:
        c = regulation_store._connection()
        c.execute(
            "UPDATE user_documents SET enabled = ? WHERE doc_id = ?",
            [enabled, doc_id],
        )
    docs = [d for d in list_documents() if d["doc_id"] == doc_id]
    return docs[0] if docs else None


def delete_document(doc_id: str) -> bool:
    with _LOCK:
        c = regulation_store._connection()
        row = c.execute(
            "SELECT 1 FROM user_documents WHERE doc_id = ?", [doc_id]
        ).fetchone()
        if not row:
            return False
        c.begin()
        try:
            c.execute("DELETE FROM document_chunks WHERE doc_id = ?", [doc_id])
            c.execute("DELETE FROM user_documents WHERE doc_id = ?", [doc_id])
            c.commit()
        except Exception:
            c.rollback()
            raise
    return True


async def retrieve_relevant_chunks(query: str, top_k: int = 5) -> list[dict[str, Any]]:
    """Семантический поиск по чанкам ENABLED-документов.

    Если эмбеддинги отсутствуют (Ollama не запущена) — fallback на keyword-match
    по lower-case substring. Возвращает [{doc_id, filename, text, score}, ...].
    """
    with _LOCK:
        c = regulation_store._connection()
        rows = c.execute(
            """
            SELECT dc.chunk_id, dc.doc_id, dc.text, dc.embedding, ud.filename
            FROM document_chunks dc
            JOIN user_documents ud ON ud.doc_id = dc.doc_id
            WHERE ud.enabled = TRUE
            """
        ).fetchall()
    if not rows:
        return []

    # Семантический путь — bge-m3
    qvec = await embed_query(query)
    if qvec is not None:
        scored: list[tuple[float, str, str, str]] = []
        for chunk_id, doc_id, text, emb_json, filename in rows:
            if not emb_json:
                continue
            try:
                vec = json.loads(emb_json)
            except (TypeError, ValueError):
                continue
            score = _cosine(qvec, vec)
            scored.append((score, doc_id, filename, text))
        scored.sort(key=lambda x: x[0], reverse=True)
        return [
            {"doc_id": d, "filename": f, "text": t, "score": round(s, 3)}
            for s, d, f, t in scored[:top_k]
        ]

    # Fallback: keyword-match (грубый, но лучше чем ничего)
    q_words = {w for w in re.findall(r"\w{3,}", query.lower()) if len(w) >= 3}
    if not q_words:
        return []
    scored2: list[tuple[int, str, str, str]] = []
    for _chunk_id, doc_id, text, _emb, filename in rows:
        lower = text.lower()
        hits = sum(1 for w in q_words if w in lower)
        if hits > 0:
            scored2.append((hits, doc_id, filename, text))
    scored2.sort(key=lambda x: x[0], reverse=True)
    return [
        {"doc_id": d, "filename": f, "text": t, "score": float(s)}
        for s, d, f, t in scored2[:top_k]
    ]
